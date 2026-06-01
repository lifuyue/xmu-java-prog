from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
REPORT_DIR = ROOT / "report"
SCREENSHOT_DIR = ROOT / "screenshots"
OUT = REPORT_DIR / "Java图书管理系统实验文档.docx"
DIAGRAM = REPORT_DIR / "class-diagram.png"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths):
            if index >= len(row.cells):
                continue
            cell = row.cells[index]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "F2F4F7")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], str(value))
        set_table_geometry(table, widths)
    return table


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def create_class_diagram():
    width, height = 1700, 920
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 38)
        box_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 25)
        small_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 21)
    except OSError:
        title_font = ImageFont.load_default()
        box_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.text((52, 36), "Library Management System Class Diagram", fill="#1F4D78", font=title_font)
    boxes = {
        "User": (80, 150, 440, 360, ["username", "password", "displayName", "role"]),
        "AdminUser": (80, 470, 440, 640, ["permissionDescription()"]),
        "ReaderUser": (80, 710, 440, 880, ["permissionDescription()"]),
        "Book": (610, 150, 970, 390, ["isbn", "title", "author", "publisher", "stock"]),
        "BorrowRecord": (610, 520, 970, 850, ["recordId", "readerName", "isbn", "bookTitle", "borrowDate", "returnDate", "returned"]),
        "LibrarySystem": (1190, 170, 1600, 530, ["books: Map", "records: List", "users: Map", "addBook()", "borrowBook()", "returnBook()", "searchBooks()"]),
        "Repository": (1190, 650, 1600, 850, ["BookRepository", "BorrowRecordRepository", "UserRepository"]),
    }
    for name, (x1, y1, x2, y2, fields) in boxes.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=12, outline="#2E74B5", width=4, fill="#F8FAFC")
        draw.rectangle((x1, y1, x2, y1 + 52), fill="#E8EEF5", outline="#2E74B5", width=2)
        draw.text((x1 + 18, y1 + 13), name, fill="#0B2545", font=box_font)
        yy = y1 + 72
        for field in fields:
            draw.text((x1 + 22, yy), field, fill="#1F2937", font=small_font)
            yy += 30

    def line_with_arrow(points, label, label_at):
        draw.line(points, fill="#4B5563", width=4, joint="curve")
        x2, y2 = points[-1]
        x1, y1 = points[-2]
        if abs(x2 - x1) > abs(y2 - y1):
            if x2 > x1:
                arrow_head = [(x2, y2), (x2 - 16, y2 - 8), (x2 - 16, y2 + 8)]
            else:
                arrow_head = [(x2, y2), (x2 + 16, y2 - 8), (x2 + 16, y2 + 8)]
        else:
            if y2 > y1:
                arrow_head = [(x2, y2), (x2 - 8, y2 - 16), (x2 + 8, y2 - 16)]
            else:
                arrow_head = [(x2, y2), (x2 - 8, y2 + 16), (x2 + 8, y2 + 16)]
        draw.polygon(arrow_head, fill="#4B5563")
        draw.rectangle((label_at[0] - 8, label_at[1] - 4, label_at[0] + 125, label_at[1] + 28), fill="white")
        draw.text(label_at, label, fill="#4B5563", font=small_font)

    line_with_arrow([(260, 470), (260, 360)], "extends", (285, 402))
    line_with_arrow([(260, 710), (260, 360)], "extends", (285, 582))
    line_with_arrow([(1190, 310), (970, 270)], "manages", (1020, 230))
    line_with_arrow([(1190, 420), (970, 685)], "records", (1020, 535))
    line_with_arrow([(1395, 650), (1395, 530)], "loads/saves", (1415, 575))
    DIAGRAM.parent.mkdir(parents=True, exist_ok=True)
    img.save(DIAGRAM)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Java 图书管理系统大作业"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def build_doc():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    create_class_diagram()
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Java 图书管理系统实验文档")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("学号：22920242203267    姓名：李富悦    日期：2026 年 6 月 1 日")

    doc.add_heading("1. 项目概述", level=1)
    add_body(doc, "本项目按照 Java 大作业要求实现图书管理系统。系统使用 JavaFX 构建桌面界面，区分管理员和普通读者两种角色，支持图书管理、图书查询、借书、还书、查看未归还记录和文本文件持久化。")
    add_body(doc, "工程包名为 com.library.lifuyue，采用 model、repository、service、ui 分层结构。数据文件保存在 data 目录，程序启动时读取，操作后和退出时保存。")

    doc.add_heading("2. 类图", level=1)
    doc.add_picture(str(DIAGRAM), width=Inches(6.3))
    add_caption(doc, "图1 系统核心类图")

    doc.add_heading("3. 核心功能说明", level=1)
    add_table(doc, ["模块", "主要类", "功能说明"], [
        ["角色登录", "AuthService, User, AdminUser, ReaderUser", "管理员使用 users.txt 中账号密码登录，普通读者输入姓名进入；User 子类体现继承和多态。"],
        ["图书管理", "Book, LibrarySystem", "管理员可添加、修改、删除、查看图书；删除前确认，存在未归还记录时禁止删除。"],
        ["图书查询", "BookTablePane, LibrarySystem", "支持按书名模糊查询、按作者查询、按 ISBN 精确查询，并可重置查看全部。"],
        ["借阅管理", "BorrowRecord, BorrowPane", "借书时库存减 1，归还时库存加 1；库存为 0 或重复借阅时给出提示。"],
        ["数据持久化", "BookRepository, BorrowRecordRepository, UserRepository", "通过 FileReader、FileWriter、BufferedReader、BufferedWriter 读写文本数据。"],
    ], [1800, 2600, 4960])

    doc.add_heading("4. 数据持久化设计", level=1)
    add_table(doc, ["文件", "保存内容", "格式"], [
        ["books.txt", "全部图书信息", "ISBN,书名,作者,出版社,库存"],
        ["records.txt", "全部借阅记录", "记录编号,读者姓名,ISBN,书名,借阅日期,归还日期,是否归还"],
        ["users.txt", "管理员和读者用户", "账号,密码,角色,显示姓名"],
    ], [1800, 2500, 5060])
    add_body(doc, "读取文件时如果某一行格式错误，系统跳过该行并继续加载其他数据，避免单条坏数据导致整个系统无法启动。")

    doc.add_heading("5. 运行截图", level=1)
    screenshot_items = [
        ("login.png", "图2 登录与角色入口"),
        ("book-management.png", "图3 管理员图书管理"),
        ("book-query.png", "图4 图书查询"),
        ("borrow-records.png", "图5 借阅记录"),
        ("reader-borrow.png", "图6 读者借阅归还"),
        ("reader-records.png", "图7 读者我的借阅"),
    ]
    for filename, caption in screenshot_items:
        path = SCREENSHOT_DIR / filename
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.4))
            add_caption(doc, caption)

    doc.add_heading("6. 测试与验证", level=1)
    add_bullet(doc, "编译验证：执行 mvn clean package，确认项目可以正常编译打包。")
    add_bullet(doc, "登录验证：管理员账号 admin/admin123 可登录，错误密码会提示失败。")
    add_bullet(doc, "查询验证：书名、作者、ISBN 三种查询方式均能过滤 TableView 中的图书。")
    add_bullet(doc, "借阅验证：库存充足时借书成功，库存为 0 和重复借阅时系统拒绝操作。")
    add_bullet(doc, "持久化验证：借书、还书、添加读者后，books.txt、records.txt、users.txt 会保存最新状态。")

    doc.add_heading("7. 遇到的问题及解决方法", level=1)
    add_table(doc, ["问题", "解决方法"], [
        ["界面、业务和文件读写容易混杂在一个类中。", "将系统拆分为 model、repository、service、ui 四层，界面只调用服务层。"],
        ["删除图书可能破坏借阅记录一致性。", "删除前检查是否存在未归还记录，有记录时禁止删除。"],
        ["文本文件读取可能遇到格式错误或文件不存在。", "仓储层创建缺失文件，读取时跳过坏行并通过异常提示读写失败。"],
        ["真实截图容易截到其他窗口。", "使用真实 JavaFX 程序的 demo 页签启动参数，逐页打开实际窗口后截图。"],
    ], [3300, 6060])

    doc.add_heading("8. 实验总结", level=1)
    add_body(doc, "本次作业通过 JavaFX、集合框架、异常处理和 IO 流综合实现一个小型图书管理系统。分层设计使核心业务规则更容易测试，也便于在报告和答辩中说明每个类的职责。")
    add_body(doc, "AI 辅助开发适合用于拆解需求、生成初稿和补充测试清单，但最终仍需要人工运行、截图、检查持久化文件和修正界面细节，才能保证提交材料符合课程要求。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
