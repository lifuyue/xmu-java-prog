from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "JAVA实验报告LAB8.docx"
SCREENSHOTS = ROOT / "screenshots"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_font(paragraph, size=11):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def clear_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=(35, 94, 173))
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.25)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("• " + text)
    set_run_font(run, size=10.5)
    return paragraph


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=(90, 90, 90))


def add_code(document, code):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(code)
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
    run.font.size = Pt(8.5)


def add_info_table(document):
    rows = [
        ("实验名称：", "LAB8 Java GUI 编程"),
        ("实验日期：", "20260425"),
        ("实验地点：", "文宣楼B308"),
        ("提交日期：", "20260425"),
        ("学号：", "22920242203267"),
        ("姓名：", "李富悦"),
        ("专业年级：", "软工2024"),
    ]
    box = document.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.space_after = Pt(10)
    for index, (left, right) in enumerate(rows):
        if index:
            box.add_run("\n")
        left_run = box.add_run(left)
        set_run_font(left_run, size=11, bold=True, color=(70, 70, 70))
        value_run = box.add_run(right)
        set_run_font(value_run, size=11)
    document.add_paragraph()


def add_screenshot(document, filename, caption, width=5.6):
    path = SCREENSHOTS / filename
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(document, caption)


def main():
    document = Document(DOCX)
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    clear_body(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("《Java程序设计》实验报告")
    set_run_font(run, size=22, bold=True, color=(35, 94, 173))

    add_info_table(document)

    add_heading(document, "实验目的")
    for item in [
        "掌握 Java Swing 图形界面编程中的常用组件、容器和布局管理器。",
        "理解事件源、事件对象和事件监听器之间的关系，能够编写按钮、列表和鼠标事件处理逻辑。",
        "将控制台输入输出程序改造为图形界面程序，理解抽象类在复用业务流程中的作用。",
        "掌握 JTable 与 TableModel 的基本使用方式，能够用表格展示和维护结构化数据。",
        "理解 MVC 模式在 GUI 程序中的分层思想，降低模型、界面和控制逻辑之间的耦合。",
    ]:
        add_bullet(document, item)

    add_heading(document, "实验内容")
    add_body(
        document,
        "本次实验依据第 08 周 Java GUI 编程要求完成。主要内容包括：运行并改写登录框界面；"
        "将 GuessGame 改造成图形界面版本；编写事件处理示例；使用 JTable 与 MVC 完成学生管理系统图形界面。"
    )
    for item in [
        "题目 1：使用 JFrame、JTextField、JPasswordField、JButton 等组件实现登录框。",
        "题目 2：设计抽象类 GuessGame，分别实现控制台版 ConsoleGuessGame 与对话框版 GUIGame。",
        "题目 3：通过按钮、下拉框和鼠标进入/离开区域演示 Java 事件处理机制。",
        "题目 5：使用 JTable 展示学生数据，并采用 Model、View、Controller 分层完成新增、修改、删除和查找。",
    ]:
        add_bullet(document, item)

    add_heading(document, "算法流程")
    add_heading(document, "1. GuessGame 图形界面流程", level=2)
    for item in [
        "抽象父类生成 1 到 100 的随机数，并控制一轮猜数循环。",
        "每次调用 readGuess 获取用户输入，图形界面版本通过 JOptionPane 弹出输入框。",
        "根据输入与答案的大小关系调用 showMessage 给出提示。",
        "猜中或取消后调用 askContinue 询问是否继续下一轮。",
    ]:
        add_bullet(document, item)
    add_heading(document, "2. 学生管理系统 MVC 流程", level=2)
    for item in [
        "StudentModel 保存学生列表，并提供新增、修改、删除、按学号查找和平均分统计方法。",
        "StudentTableModel 将 StudentModel 中的数据适配给 JTable 显示。",
        "StudentView 只负责界面组件和表单控件，不直接修改数据。",
        "StudentController 监听按钮和表格选择事件，读取表单、调用模型并刷新表格。",
    ]:
        add_bullet(document, item)

    add_heading(document, "核心功能实现")
    add_body(
        document,
        "第 1 题使用 BorderLayout 组织标题、表单和按钮区，登录按钮通过 ActionListener 校验用户名与密码，"
        "清空按钮负责重置文本框并更新状态提示。"
    )
    add_code(
        document,
        "loginButton.addActionListener(event -> {\n"
        "    String name = userField.getText().trim();\n"
        "    if (name.isEmpty() || passwordField.getPassword().length == 0) {\n"
        "        status.setText(\"用户名和密码不能为空。\");\n"
        "    } else {\n"
        "        status.setText(\"欢迎，\" + name + \"。登录信息已提交。\");\n"
        "    }\n"
        "});"
    )
    add_body(
        document,
        "第 2 题的关键是把游戏流程写在抽象类中，控制台版和图形界面版只实现输入、输出和继续确认三个抽象方法。"
        "这样同一套猜数字规则可以复用于不同界面。"
    )
    add_code(
        document,
        "protected abstract Integer readGuess(String prompt);\n"
        "protected abstract void showMessage(String message);\n"
        "protected abstract boolean askContinue();"
    )
    add_body(
        document,
        "第 5 题采用 MVC 结构。控制器接收按钮事件后读取表单数据，调用模型进行业务校验与数据更新，"
        "再通知 StudentTableModel 刷新 JTable。"
    )

    add_heading(document, "运行结果截图")
    add_screenshot(document, "lab8-1-login.png", "图 1  登录框界面运行效果", width=5.2)
    add_screenshot(document, "lab8-2-guessgame.png", "图 2  GUIGame 图形界面运行效果", width=5.2)
    add_screenshot(document, "lab8-3-event.png", "图 3  事件处理示例运行效果", width=5.4)
    add_screenshot(document, "lab8-5-student-mvc.png", "图 4  JTable + MVC 学生管理系统运行效果", width=5.8)

    add_heading(document, "实验结果与分析")
    add_body(
        document,
        "四个题目均已在 LAB8 目录下按题号建立独立工程目录，并通过 javac 编译。"
        "登录框能够完成输入、清空和登录反馈；GuessGame 同时保留控制台版与图形界面版；"
        "事件处理示例能够响应按钮、下拉框和鼠标事件；学生管理系统能够通过 JTable 完成数据展示和基本维护。"
    )
    add_body(
        document,
        "本次实验中，MVC 分层对 GUI 程序的结构帮助最明显。若把表格刷新、数据校验和按钮响应全部写在窗口类中，"
        "代码会迅速变得臃肿；拆分后，StudentModel 负责数据规则，StudentView 负责界面展示，StudentController 负责协调，"
        "程序更容易调试和扩展。"
    )

    add_heading(document, "实验总结")
    add_body(
        document,
        "通过本次实验，我熟悉了 Swing 中常用组件和布局管理器的基本使用方法，并进一步理解了事件处理机制。"
        "在 GuessGame 中，抽象类保留稳定流程、子类替换输入输出方式，体现了面向对象复用。"
        "在学生管理系统中，JTable 与 MVC 的组合使界面展示和业务数据分离，提升了程序的可维护性。"
    )

    document.save(DOCX)


if __name__ == "__main__":
    main()
