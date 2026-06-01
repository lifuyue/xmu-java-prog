from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path("/Users/lifuyue/Documents/model/JAVA实验报告.docx")
OUT = ROOT / "report" / "JAVA大作业1-AI记事本实验报告.docx"
SCREENSHOTS = ROOT / "screenshots"


def _children(element):
    return list(element)


def _clear_paragraph_keep_props(paragraph):
    p = paragraph._p
    for child in _children(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _template_run_props(p_xml):
    p_pr = p_xml.find(qn("w:pPr"))
    if p_pr is not None:
        r_pr = p_pr.find(qn("w:rPr"))
        if r_pr is not None:
            return deepcopy(r_pr)

    for run in p_xml.findall(qn("w:r")):
        r_pr = run.find(qn("w:rPr"))
        if r_pr is not None:
            return deepcopy(r_pr)
    return None


def _append_text_run(p_xml, text, r_pr=None):
    run = OxmlElement("w:r")
    if r_pr is not None:
        run.append(deepcopy(r_pr))

    text_el = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    run.append(text_el)
    p_xml.append(run)


def _set_alignment(p_xml, value):
    p_pr = p_xml.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_xml.insert(0, p_pr)

    for child in list(p_pr):
        if child.tag == qn("w:jc"):
            p_pr.remove(child)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), value)
    p_pr.append(jc)


def _clone_paragraph(template_p_xml, text="", align=None, keep_next=False):
    p_xml = deepcopy(template_p_xml)
    r_pr = _template_run_props(p_xml)

    for child in _children(p_xml):
        if child.tag != qn("w:pPr"):
            p_xml.remove(child)

    if text:
        _append_text_run(p_xml, text, r_pr)

    if align is not None:
        _set_alignment(p_xml, align)

    if keep_next:
        p_pr = p_xml.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            p_xml.insert(0, p_pr)
        if p_pr.find(qn("w:keepNext")) is None:
            p_pr.append(OxmlElement("w:keepNext"))

    return p_xml


def _append_before_section(doc, element):
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(element)
    else:
        body.insert(body.index(sect_pr), element)
    return Paragraph(element, doc._body)


def _trim_body_after_cover_table(doc):
    body = doc._body._element
    cover_table = doc.tables[0]._tbl
    cover_index = body.index(cover_table)

    for element in list(body)[cover_index + 1 :]:
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def _set_paragraph_text_from_template(paragraph, text):
    p_xml = paragraph._p
    r_pr = _template_run_props(p_xml)
    _clear_paragraph_keep_props(paragraph)
    _append_text_run(p_xml, text, r_pr)


def _fill_cover_table(doc):
    values = {
        0: "AI 编程作业：Swing 记事本程序",
        1: "2026 年 5 月 28 日",
        2: "个人电脑 / Codex 开发环境",
        3: "2026 年 5 月 28 日",
        5: "22920242203267",
        6: "李富悦",
        7: "2024 级",
        8: "2025-2026 学年第二学期",
    }

    table = doc.tables[0]
    for row_index, value in values.items():
        paragraph = table.rows[row_index].cells[1].paragraphs[0]
        _set_paragraph_text_from_template(paragraph, value)


def _add_heading(doc, template, text):
    _append_before_section(doc, _clone_paragraph(template, text))


def _add_body(doc, template, text="", align=None, keep_next=False):
    return _append_before_section(
        doc, _clone_paragraph(template, text, align=align, keep_next=keep_next)
    )


def _add_lines(doc, template, lines):
    for line in lines:
        _add_body(doc, template, line)


def _add_image(doc, body_template, image_name, caption, width):
    path = SCREENSHOTS / image_name
    if not path.exists():
        _add_body(doc, body_template, f"{caption}：截图文件缺失（{image_name}）。")
        return

    image_p = _append_before_section(
        doc, _clone_paragraph(body_template, "", align="center", keep_next=True)
    )
    image_p.add_run().add_picture(str(path), width=Inches(width))
    _add_body(doc, body_template, caption, align="center")


def _add_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document(TEMPLATE)

    heading_indexes = [11, 12, 13, 15, 18, 19, 21, 22]
    heading_templates = [deepcopy(doc.paragraphs[i]._p) for i in heading_indexes]
    body_template = deepcopy(doc.paragraphs[14]._p)

    _fill_cover_table(doc)
    _trim_body_after_cover_table(doc)
    _add_page_break(doc)

    _add_heading(doc, heading_templates[0], "实验目的 ")
    _add_lines(
        doc,
        body_template,
        [
            "本次实验要求借助 AI 工具完成一个 Java 记事本程序，功能尽量接近 Windows 记事本，并整理代码、实验报告和演示视频材料。",
            "实验重点是把自然语言需求拆解为可实现的 Swing 桌面程序功能，使用 AI 辅助完成代码设计、调试、截图材料和报告撰写，同时保留人工检查与修改过程。",
        ],
    )

    _add_heading(doc, heading_templates[1], "实验内容")
    _add_lines(
        doc,
        body_template,
        [
            "1. 使用 Java Swing 实现桌面记事本程序，包含新建、打开、保存、另存为、退出、撤销、剪切、复制、粘贴、删除、查找、替换、转到、自动换行、字体、缩放、状态栏、关于等常用功能。",
            "2. 按课程要求保存源代码、实验报告和程序运行视频。录制视频时，在记事本中输入学号和姓名：22920242203267 李富悦。",
            "3. 使用 AI 辅助生成代码、解释功能实现、定位编译问题和整理报告材料，并将关键提示词和 AI 生成代码过程截图放入报告。",
            "主要提示词：请使用 Java Swing 开发一个接近 Windows 记事本的桌面程序，要求有菜单栏、文本编辑区、文件读写、编辑功能、查找替换、字体设置、自动换行、状态栏和缩放功能，代码结构清晰，适合作为 Java 大作业提交。",
            "补充提示词：请根据课程实验报告模板，生成一份包含实验目的、实验内容、算法流程、核心功能实现、测试数据、运行结果、问题解决和实验心得的 Word 报告，并加入 AI 生成代码过程截图。",
        ],
    )

    _add_heading(doc, heading_templates[2], "算法流程（根据需要选择）")
    _add_lines(
        doc,
        body_template,
        [
            "1. 程序启动时创建 NotepadFrame 主窗口，初始化菜单栏、文本编辑区、状态栏和快捷键，并创建 DocumentModel 保存当前文件、文本修改状态和编码信息。",
            "2. 用户执行新建、打开、保存、另存为等操作时，先检查当前文档是否已经修改；如果未保存，则弹出确认对话框，再调用 FileService 完成 UTF-8 文本读写。",
            "3. 文本编辑区通过 Swing 文档监听器记录修改状态，通过 UndoManager 支持撤销和重做，通过 JTextArea 自带方法支持剪切、复制、粘贴和删除。",
            "4. 查找和替换功能由 FindReplaceDialog 处理，按用户输入的关键词在当前文本中定位匹配位置，并可执行单次替换或全部替换。",
            "5. 字体、自动换行、缩放和状态栏属于界面状态控制，修改后立即作用到 JTextArea，并同步窗口底部的行列号、字符数和缩放比例。",
        ],
    )

    _add_heading(doc, heading_templates[3], "核心功能实现")
    _add_lines(
        doc,
        body_template,
        [
            "NotepadApp 是程序入口，负责设置系统观感并启动主窗口。NotepadFrame 负责菜单、工具动作、文本区、状态栏和窗口关闭逻辑，是整个程序的界面控制中心。",
            "DocumentModel 负责记录当前文件路径、修改状态和最后保存内容，避免界面层直接管理过多文档状态。FileService 统一处理打开、保存和另存为，便于后续扩展编码或异常处理。",
            "FindReplaceDialog 实现查找、查找下一个、替换和全部替换。FontChooserDialog 提供字体、字形和字号选择，使文本区可以像 Windows 记事本一样切换显示字体。",
            "StatusBar 根据光标位置和文档内容实时显示行列号、字符数、编码和缩放比例。ScreenshotExporter 用于生成报告中需要的程序运行截图，便于复现实验材料。",
        ],
    )

    _add_heading(doc, heading_templates[4], "测试数据")
    _add_lines(
        doc,
        body_template,
        [
            "测试输入文本：22920242203267 李富悦；Java Swing Notepad 测试；查找关键字 notepad；替换关键字 记事本。",
            "文件功能测试：新建文档后输入文本，执行保存；重新打开文件检查内容是否一致；修改内容后执行另存为；关闭窗口时检查未保存提示是否出现。",
            "编辑功能测试：输入多行文本后执行复制、粘贴、剪切、删除、撤销和重做；使用查找替换对多处关键词进行定位和批量替换。",
            "显示功能测试：切换自动换行，调整字体和字号，执行放大、缩小和恢复默认缩放，并观察状态栏中行列号、字符数和缩放比例是否同步变化。",
        ],
    )

    _add_heading(doc, heading_templates[5], "运行结果")
    _add_image(doc, body_template, "ai-process-1.png", "图1 AI辅助编程需求分析截图", 5.2)
    _add_image(doc, body_template, "ai-process-2.png", "图2 AI辅助编程提示词选择截图", 5.2)
    _add_image(doc, body_template, "ai-process-3.png", "图3 AI辅助编程设计讨论截图", 5.2)
    _add_image(doc, body_template, "ai-process-4.png", "图4 AI辅助生成代码过程截图", 5.2)
    _add_image(doc, body_template, "compile-run.png", "图5 程序编译与运行截图", 5.2)
    _add_image(doc, body_template, "main-window.png", "图6 记事本主界面截图", 5.2)
    _add_image(doc, body_template, "find-replace.png", "图7 查找替换功能截图", 4.6)
    _add_image(doc, body_template, "font-dialog.png", "图8 字体设置功能截图", 4.6)
    _add_image(doc, body_template, "zoom-status.png", "图9 缩放与状态栏截图", 5.2)

    _add_heading(doc, heading_templates[6], "出现的问题及解决方法")
    _add_lines(
        doc,
        body_template,
        [
            "问题一：AI 初稿容易把所有功能写在一个类中，后续维护困难。解决方法是继续提示 AI 按职责拆分为主窗口、文档模型、文件服务、查找替换对话框和字体对话框。",
            "问题二：保存和退出时容易遗漏未保存确认。解决方法是在新建、打开、退出前统一调用确认方法，保证用户修改内容不会被直接覆盖。",
            "问题三：查找替换需要处理大小写、当前位置和未找到提示。解决方法是把查找逻辑集中在对话框中，并对匹配失败、空关键词等情况进行判断。",
            "问题四：报告排版若直接指定字体，容易偏离学校模板。解决方法是保留原始 Word 模板的段落、编号和表格格式，只替换文字和插入图片。",
        ],
    )

    _add_heading(doc, heading_templates[7], "实验心得")
    _add_lines(
        doc,
        body_template,
        [
            "AI 在快速搭建 Swing 程序结构、生成菜单动作、补齐异常处理和整理测试步骤方面效果较好，能节省大量重复编码时间。对于明确的小功能，例如文件读写、查找替换和状态栏更新，提示词越具体，生成结果越接近可直接使用的代码。",
            "使用 AI 时不能只给一句笼统要求。更有效的提示词应说明开发语言、界面框架、目标功能、文件结构、交互细节和提交要求；生成后还需要人工编译运行、检查边界情况，并根据课程模板调整报告格式。",
            "本次实验中遇到的主要困难是功能较多，AI 初稿容易忽略细节或产生不符合模板的排版。通过分步骤提示、逐项测试和对照原始模板修改，最终完成了接近 Windows 记事本的 Java 程序和配套实验报告。",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
