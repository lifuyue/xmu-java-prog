from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path("/Users/lifuyue/Documents/hw/java/JAVA实验报告.docx")
DOCX = ROOT / "JAVA实验报告LAB9.docx"
SCREENSHOTS = ROOT / "screenshots"


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def fill_original_info_table(document):
    table = document.tables[0]
    replacements = {
        0: "LAB9",
        1: "20260511",
        2: "文宣楼B308",
        3: "20260503",
        5: "22920242203267",
        6: "李富悦",
        7: "软工2024",
        8: "大二下",
    }
    for row_index, value in replacements.items():
        set_cell_text(table.rows[row_index].cells[1], value)


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Cannot find paragraph: {exact_text}")


def clone_paragraph_after(paragraph, text="", style_source=None):
    source = style_source or paragraph
    new_p = deepcopy(source._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_paragraph.add_run(text)
        if source.runs:
            run._rPr = deepcopy(source.runs[0]._r.rPr)
        run.font.size = Pt(10.5)
    return new_paragraph


def insert_text_block_after(anchor, lines, normal_source):
    current = anchor
    for line in reversed(lines):
        current = clone_paragraph_after(anchor, line, normal_source)
    return current


def add_picture_after(anchor, image_path, caption, normal_source, width=5.8):
    caption_paragraph = clone_paragraph_after(anchor, caption, normal_source)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    picture_paragraph = clone_paragraph_after(anchor, "", normal_source)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    return caption_paragraph


def remove_original_note_if_needed(note):
    # 保留原模板的“说明：最好截图说明”，只把它作为运行结果部分的提示行。
    note.paragraph_format.space_after = Pt(4)


def fill_report_body(document):
    normal_source = find_paragraph(document, "说明：最好截图说明")

    sections = {
        "实验目的": [
            "1. 熟悉 JavaFX 程序结构，掌握 Application、Stage、Scene 与常用布局容器的使用。",
            "2. 掌握 TableView、ObservableList、FilteredList、TextField、ComboBox、Button 等组件的组合方式。",
            "3. 在给定 HelloApplication.java 的学生查询界面基础上补完整课程管理和选课管理界面。",
            "4. 掌握 Canvas 绘图和鼠标事件处理，能够通过拖拽绘制矩形、圆形和直线。",
            "5. 学习使用颜色选择对话框修改画笔颜色，并将颜色应用到后续绘制图形。",
        ],
        "实验内容": [
            "本次实验完成 lab9.docx 中的两个基本题。第一题是在学生管理系统示例基础上继续扩展课程管理和选课管理界面，系统能够展示学生、课程和选课记录，并支持课程新增、修改、删除以及选课记录维护。",
            "第二题是在 Painter 示例思路上实现 JavaFX 绘图面板，提供矩形、圆形、直线三种绘图工具，并通过颜色选择对话框修改图形描边颜色。",
        ],
        "算法流程（根据需要选择）": [
            "1. 程序启动后创建 TabPane，其中“学生选课管理”标签页放置左侧菜单和右侧工作区，“Painter绘图”标签页放置绘图面板。",
            "2. 点击学生管理、课程管理、选课管理按钮时，右侧工作区分别切换到查询、课程维护和选课维护界面。",
            "3. 学生查询使用 FilteredList 根据学号、电话、班级和政治面貌筛选 TableView 数据。",
            "4. 课程管理通过表单采集课程号、课程名称、学分、教师、课程类型和上课时间，完成新增、修改和删除。",
            "5. 选课管理通过学生 ComboBox 和课程 ComboBox 建立关联，新增、修改、删除后刷新选课表和统计信息。",
            "6. Painter 在鼠标按下时记录起点，拖拽时绘制预览图形，释放鼠标后把图形保存到列表并重绘画布。",
        ],
        "核心功能实现": [
            "1. 课程管理界面使用 ObservableList<Course> 保存课程数据，TableView 负责展示课程号、课程名称、学分、任课教师、课程类型和上课时间。新增课程前检查课程号是否重复，删除课程时同步移除相关选课记录。",
            "2. 选课管理界面使用 Enrollment 保存学生与课程对象引用，新增前检查同一学生是否已经选择同一门课程，成绩输入为空或不合法时按“未录入”处理。",
            "3. PainterPane 使用 Canvas 绘图，并用 PaintShape 记录图形类型、起点、终点和颜色。每次拖拽预览、撤销和清空后统一调用 drawAll 方法重绘，避免画布残影。",
            "4. 颜色修改使用 Dialog<Color> 和 ColorPicker，点击“确定”后更新 currentColor，后续绘制的矩形、圆形和直线均使用新颜色。",
        ],
        "测试数据": [
            "1. 学生数据：内置 7 条学生记录，包含学号、姓名、电话、班级、政治面貌和邮箱。",
            "2. 课程数据：内置 Java程序设计、数据库系统、人机交互设计 3 门课程。",
            "3. 选课数据：内置 4 条选课记录，包含已结课、已选和未录入成绩等情况。",
            "4. 测试操作：按班级查询学生；新增 AI401 人工智能基础课程；新增学生选课记录；在 Painter 中分别绘制矩形、圆形、直线并修改颜色。",
            "5. 测试结果：各模块均能按预期响应，表格数据和绘图结果显示正常。",
        ],
        "出现的问题及解决方法": [
            "1. 问题：JavaFX 不再随 JDK 默认内置，直接编译会缺少 javafx.controls 模块。解决方法：建立 Maven 项目，在 pom.xml 中引入 org.openjfx:javafx-controls，并使用 javafx-maven-plugin 运行主类。",
            "2. 问题：课程对象修改后，选课记录中引用的课程信息需要同步显示。解决方法：修改课程后调用 TableView.refresh，并同步刷新选课表。",
            "3. 问题：Canvas 拖拽预览如果直接连续绘制，会出现预览残影。解决方法：维护图形列表，每次拖拽先清空画布、重绘历史图形，再绘制当前预览。",
        ],
        "实验心得": [
            "通过本次实验，我进一步熟悉了 JavaFX 表格、表单、布局容器和事件处理的组合方式。课程管理和选课管理的补全让我理解了界面切换、数据集合和表格刷新之间的关系；Painter 绘图部分让我认识到 Canvas 程序需要自己维护图形状态，撤销、清空和预览都依赖稳定的数据结构。",
        ],
    }

    for heading, lines in sections.items():
        anchor = find_paragraph(document, heading)
        insert_text_block_after(anchor, lines, normal_source)

    note = find_paragraph(document, "说明：最好截图说明")
    remove_original_note_if_needed(note)
    current = note
    pictures = [
        ("lab9-student-management.png", "图1 学生管理查询界面"),
        ("lab9-course-management.png", "图2 课程管理界面"),
        ("lab9-enrollment-management.png", "图3 选课管理界面"),
        ("lab9-painter.png", "图4 Painter绘图界面"),
    ]
    for filename, caption in pictures:
        current = add_picture_after(current, SCREENSHOTS / filename, caption, normal_source)


def fix_schema_without_changing_layout(document):
    section = document.sections[0]
    section._sectPr.pgMar.set(qn("w:gutter"), "0")

    settings = document.settings._element
    for item in list(settings.findall(qn("w:displayBackgroundShape"))):
        settings.remove(item)


def main():
    document = Document(TEMPLATE)
    document.core_properties.author = "李富悦"
    document.core_properties.modified = datetime(2026, 5, 3, 23, 55)

    fill_original_info_table(document)
    fill_report_body(document)
    fix_schema_without_changing_layout(document)

    document.save(DOCX)


if __name__ == "__main__":
    main()
