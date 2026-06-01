from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path("/Users/lifuyue/Documents/model/JAVA实验报告.docx")
DOCX = ROOT / "JAVA实验报告LAB10.docx"
SCREENSHOTS = ROOT / "screenshots"


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def fill_info_table(document):
    table = document.tables[0]
    values = {
        0: "LAB10 文件处理",
        1: "2026年5月28日",
        2: "文宣楼B308",
        3: "2026年5月28日",
        5: "22920242203267",
        6: "李富悦",
        7: "软件工程2024级",
        8: "大二下",
    }
    for row_index, value in values.items():
        set_cell_text(table.rows[row_index].cells[1], value)


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Cannot find paragraph: {exact_text}")


def clone_paragraph_after(anchor, text="", style_source=None):
    source = style_source or anchor
    new_p = deepcopy(source._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if text:
        run = paragraph.add_run(text)
        if source.runs:
            run._rPr = deepcopy(source.runs[0]._r.rPr)
        run.font.size = Pt(10.5)
    return paragraph


def insert_lines_after(anchor, lines, normal_source):
    current = anchor
    for line in reversed(lines):
        current = clone_paragraph_after(anchor, line, normal_source)
    return current


def add_picture_after(anchor, image_path, caption, normal_source, width=5.8):
    caption_paragraph = clone_paragraph_after(anchor, caption, normal_source)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    picture_paragraph = clone_paragraph_after(anchor, "", normal_source)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    return caption_paragraph


def fill_body(document):
    normal_source = find_paragraph(document, "说明：最好截图说明")

    sections = {
        "实验目的": [
            "1. 熟悉 Java 文件处理相关类的使用，掌握 Path、Files、BasicFileAttributes 等 API 对文件和文件夹进行解析的方法。",
            "2. 掌握顺序存取文件的基本思想，能够通过文本文件保存、读取、重写结构化学生信息。",
            "3. 熟悉 JavaFX 图形用户界面的设计方式，能够组合 TabPane、TableView、TextField、ImageView、FileChooser 等组件完成交互。",
            "4. 掌握学生信息新增、删除、修改、按姓名查询和显示全部等功能的事件处理流程。",
            "5. 掌握照片文件选择、格式检查和集中复制保存的方法。",
        ],
        "实验内容": [
            "本次实验完成两个基本题。第一题实现文件路径解析：用户输入文件或文件夹路径后，程序判断路径类型；若是文件夹，则统计其直接包含的文件个数和文件夹个数；若是文件，则显示文件大小和最后修改日期。",
            "第二题设计 JavaFX 图形界面，用顺序存取文本文件 data/students.txt 保存学生信息。界面允许输入学号、姓名、电话、邮箱，选择 JPG/JPEG 照片并复制到 data/photos/，同时提供新增、删除、修改、查询、显示全部以及上一条/下一条浏览功能。",
        ],
        "算法流程（根据需要选择）": [
            "1. 程序启动后创建 TabPane，分别放置“路径解析”和“学生顺序文件管理”两个标签页。",
            "2. 路径解析功能读取输入框中的路径，使用 Files.isDirectory 和 Files.isRegularFile 判断类型；文件夹通过 Files.list 统计直接子文件和子文件夹，普通文件通过 BasicFileAttributes 获取大小和最后修改时间。",
            "3. 学生管理功能启动时创建 data/ 和 data/photos/，并顺序读取 students.txt 的每一行，将制表符分隔的字段转换为 Student 对象。",
            "4. 新增学生时先校验必填字段、电话、邮箱和照片格式，再把照片复制到 data/photos/，最后把新对象加入列表并顺序写回 students.txt。",
            "5. 修改和删除都以当前正在显示的学生记录为操作对象，修改后或删除后重新按顺序保存整个文本文件。",
            "6. 查询功能按照姓名关键字过滤学生列表，显示全部功能读取当前文件中的所有学生；结果集合通过 currentIndex 控制上一条和下一条浏览。",
        ],
        "核心功能实现": [
            "1. StudentRepository 负责顺序文件读写。load 方法逐行读取 data/students.txt，save 方法把当前学生列表逐行写入临时文件，再原子替换原文件，体现顺序存取文件的保存过程。",
            "2. copyPhoto 方法检查照片后缀是否为 .jpg 或 .jpeg，把照片复制到 data/photos/，并在学生记录中保存相对路径，便于项目文件夹移动后继续运行。",
            "3. FileAnalyzerPane 封装路径解析界面，分别处理文件夹统计和普通文件属性展示，结果写入只读 TextArea。",
            "4. StudentManagerPane 封装学生管理界面，TableView 展示当前查询或显示结果，表单区域显示当前记录，ImageView 展示当前学生照片。",
            "5. ScreenshotSession 仅用于实验报告取证：在真实 JavaFX 窗口中切换功能状态，并调用 macOS 系统 screencapture 工具生成窗口截图，避免使用手工绘制或伪造截图。",
        ],
        "测试数据": [
            "1. 文件路径解析测试：选择 LAB10/data 文件夹，程序正确统计其中的 students.txt 文件和 photos 文件夹。",
            "2. 新增测试：新增学号 22920242203267、姓名李富悦、电话 13900001010、邮箱 lifuyue@example.com，并选择 JPG 照片。",
            "3. 修改测试：将李富悦的电话修改为 13900002020，邮箱修改为 lifuyue-lab10@example.com，保存后表格同步刷新。",
            "4. 查询测试：输入姓名关键字“李”，系统只显示李富悦一条记录，并支持上一条、下一条浏览。",
            "5. 显示测试：显示文件中保存的全部学生信息，当前共有张三、王芳、李富悦三条记录。",
            "6. 删除测试：构造一条临时删除记录并执行删除，程序将其从顺序文件中移除，剩余记录继续正常显示。",
        ],
        "出现的问题及解决方法": [
            "1. 问题：JavaFX 不随 JDK 默认打包，直接 javac 编译会缺少 javafx.controls。解决方法：使用 Maven 项目，在 pom.xml 中引入 org.openjfx:javafx-controls，并配置 javafx-maven-plugin 运行主类。",
            "2. 问题：顺序文件不适合直接就地删除或修改某一行。解决方法：将文件内容读入集合，完成新增、删除或修改后整体顺序写回 students.txt。",
            "3. 问题：照片如果只保存原始绝对路径，项目移动后可能无法加载。解决方法：选择照片后复制到 data/photos/，学生记录保存相对路径。",
            "4. 问题：连续自动截图时界面可能尚未完成刷新。解决方法：截图流程在每次功能状态切换后延时等待 JavaFX 重绘，再调用 screencapture。",
        ],
        "实验心得": [
            "通过本次实验，我进一步理解了顺序文件处理的特点：新增比较直接，而删除和修改通常需要重新生成文件内容。JavaFX 界面部分让我熟悉了表格、表单、文件选择器和图片预览的组合方式，也认识到界面状态与底层文件数据必须保持同步。真实窗口截图流程则帮助我确认程序运行效果和实验报告中的运行结果一致。",
        ],
    }

    for heading, lines in sections.items():
        anchor = find_paragraph(document, heading)
        insert_lines_after(anchor, lines, normal_source)

    note = find_paragraph(document, "说明：最好截图说明")
    note.paragraph_format.space_after = Pt(4)
    current = note
    screenshots = [
        ("lab10-file-analysis.png", "图1 本机 JavaFX 窗口截图：文件夹路径解析"),
        ("lab10-add-student.png", "图2 本机 JavaFX 窗口截图：新增学生信息并保存照片"),
        ("lab10-modify-student.png", "图3 本机 JavaFX 窗口截图：修改当前学生信息"),
        ("lab10-query-student.png", "图4 本机 JavaFX 窗口截图：按姓名查询学生信息"),
        ("lab10-display-students.png", "图5 本机 JavaFX 窗口截图：显示全部学生信息"),
        ("lab10-delete-student.png", "图6 本机 JavaFX 窗口截图：删除当前学生信息"),
    ]
    for filename, caption in screenshots:
        current = add_picture_after(current, SCREENSHOTS / filename, caption, normal_source)


def fix_schema_without_changing_layout(document):
    section = document.sections[0]
    section._sectPr.pgMar.set(qn("w:gutter"), "0")

    settings = document.settings._element
    for item in list(settings.findall(qn("w:displayBackgroundShape"))):
        settings.remove(item)


def main():
    document = Document(TEMPLATE)
    document.core_properties.author = "李富悦"
    document.core_properties.modified = datetime(2026, 5, 28, 12, 45)

    fill_info_table(document)
    fill_body(document)
    fix_schema_without_changing_layout(document)

    document.save(DOCX)


if __name__ == "__main__":
    main()
